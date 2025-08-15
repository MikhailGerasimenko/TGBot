import logging
from aiogram import Bot, types, Dispatcher
from aiogram.filters import Command
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove, InlineKeyboardMarkup, InlineKeyboardButton
import os
import asyncio
from typing import List, Dict, Optional, Set
from datetime import datetime, timedelta
from database import (verify_employee, log_registration_attempt, get_registration_attempts, get_all_employees,
                     log_qa_session, save_feedback, log_unanswered_question, get_analytics_stats, get_popular_questions)
from llm_client import LLMClient
import numpy as np
from docx import Document
import re
from collections import defaultdict
from config import API_TOKEN, ADMIN_CHAT_ID, DOCS_DIR as DOCUMENTS_DIR, LOGS_DIR, ONEC_EXPORT_PATH
from onec_sync import load_employees_from_file
import time

# Конфигурация для LLM (через сервис)
MAX_LENGTH = 2048
MAX_NEW_TOKENS = 512

# Глобальные переменные для клиентов
llm_client = LLMClient()

auth_logger = logging.getLogger(__name__)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(LOGS_DIR, 'bot.log')),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Инициализация бота
bot = Bot(token=API_TOKEN)

# Клавиатуры
main_kb = ReplyKeyboardMarkup(
    keyboard=[
        [KeyboardButton(text='Помощь'), KeyboardButton(text='Спросить')],
    ],
    resize_keyboard=True
)
retry_kb = ReplyKeyboardMarkup(
    keyboard=[[KeyboardButton(text='повтор')], [KeyboardButton(text='Отмена')]],
    resize_keyboard=True
)

# Вспомогательные: отправка администратору
async def send_admin_notification(text: str) -> bool:
    try:
        await bot.send_message(ADMIN_CHAT_ID, text, parse_mode='HTML')
        return True
    except Exception as e:
        logger.error(f"Ошибка отправки уведомления администратору: {e}")
        return False

# Нормализация ввода и данных
DASHES = "\u2010\u2011\u2012\u2013\u2014\u2212\uFE58\uFE63\uFF0D"
DASH_TRANS = str.maketrans({c: '-' for c in DASHES})

def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())

def normalize_name(text: str) -> str:
    return normalize_spaces(text).lower()

def normalize_employee_id(text: str) -> str:
    t = (text or "").translate(DASH_TRANS)
    t = t.replace(' ', '')  # убираем случайные пробелы внутри
    return t.upper()

# ========== ДАННЫЕ И СОСТОЯНИЯ (без Redis) ==========

# Кэш сотрудников в памяти (employee_id -> dict)
EMPLOYEES_CACHE: Dict[str, dict] = {}
# Индекс по нормализованному ID
EMPLOYEES_BY_NORM_ID: Dict[str, dict] = {}

# Попытки регистрации и состояния пользователей
REGISTRATION_ATTEMPTS: Dict[int, Dict[str, any]] = {}
MAX_ATTEMPTS = 3
RETRY_COMMANDS = ['retry', 'повтор', 'заново']
USER_STATES: Dict[int, Dict[str, any]] = {}
AUTHORIZED_USERS: Set[int] = set()
AUTHORIZED_INFO: Dict[int, Dict[str, str]] = {}

async def set_user_state(user_id: int, key: str, value: str):
    state = USER_STATES.get(user_id, {})
    state[key] = value
    USER_STATES[user_id] = state

async def get_user_state(user_id: int, key: str) -> Optional[str]:
    return USER_STATES.get(user_id, {}).get(key)

async def clear_user_state(user_id: int):
    if user_id in USER_STATES:
        del USER_STATES[user_id]

async def can_try_registration(user_id: int) -> bool:
    now_key = datetime.now().strftime('%Y%m%d')
    data = REGISTRATION_ATTEMPTS.get(user_id)
    if not data or data.get('date') != now_key:
        return True
    return data.get('count', 0) < MAX_ATTEMPTS

async def inc_registration_attempt(user_id: int):
    now_key = datetime.now().strftime('%Y%m%d')
    data = REGISTRATION_ATTEMPTS.get(user_id)
    if not data or data.get('date') != now_key:
        REGISTRATION_ATTEMPTS[user_id] = {'date': now_key, 'count': 1}
    else:
        data['count'] = data.get('count', 0) + 1
        REGISTRATION_ATTEMPTS[user_id] = data

async def get_next_attempt_time(user_id: int) -> str:
    data = REGISTRATION_ATTEMPTS.get(user_id)
    if not data:
        return "сейчас"
    next_day = datetime.strptime(data['date'], '%Y%m%d') + timedelta(days=1)
    return next_day.strftime('%d.%m.%Y в %H:%M')

async def load_employees_from_file_to_cache() -> int:
    if not ONEC_EXPORT_PATH:
        return 0
    try:
        rows = await load_employees_from_file(ONEC_EXPORT_PATH)
        EMPLOYEES_CACHE.clear()
        EMPLOYEES_BY_NORM_ID.clear()
        for emp in rows:
            emp_id = emp['employee_id']
            EMPLOYEES_CACHE[emp_id] = emp
            norm = normalize_employee_id(emp_id)
            emp['norm_id'] = norm
            emp['norm_name'] = normalize_name(emp.get('full_name', ''))
            EMPLOYEES_BY_NORM_ID[norm] = emp
        logger.info(f"Импортировано из файла: {len(rows)} записей")
        return len(rows)
    except Exception as e:
        logger.error(f"Ошибка загрузки сотрудников из файла: {e}")
        return 0

IC_CACHE = {
    "last_sync": None,
    "sync_in_progress": False
}

async def sync_with_1c():
    if IC_CACHE["sync_in_progress"]:
        return False
    try:
        IC_CACHE["sync_in_progress"] = True
        success = False
        if ONEC_EXPORT_PATH:
            cnt = await load_employees_from_file_to_cache()
            success = cnt > 0
        else:
            employees = await get_all_employees()
            EMPLOYEES_CACHE.clear()
            EMPLOYEES_BY_NORM_ID.clear()
            for emp in employees:
                EMPLOYEES_CACHE[emp["employee_id"]] = emp
                norm = normalize_employee_id(emp["employee_id"])
                emp['norm_id'] = norm
                emp['norm_name'] = normalize_name(emp.get('full_name', ''))
                EMPLOYEES_BY_NORM_ID[norm] = emp
            success = len(EMPLOYEES_CACHE) > 0
        if success:
            IC_CACHE["last_sync"] = datetime.now()
            logger.info("Successfully synced employees (memory cache)")
            return True
        return False
    except Exception as e:
        logger.error(f"Error syncing employees: {e}")
        return False
    finally:
        IC_CACHE["sync_in_progress"] = False

async def periodic_sync():
    while True:
        try:
            if (not IC_CACHE["last_sync"] or 
                datetime.now() - IC_CACHE["last_sync"] > timedelta(hours=1)):
                logger.info("Запуск периодической синхронизации")
                success = await sync_with_1c()
                if success:
                    await send_admin_notification(
                        f"✅ Синхронизация выполнена\n"
                        f"🕒 Время: {IC_CACHE['last_sync'].strftime('%Y-%m-%d %H:%M:%S')}"
                    )
                    await asyncio.sleep(3600)
                else:
                    await asyncio.sleep(300)
            else:
                await asyncio.sleep(60)
        except Exception as e:
            logger.error(f"Ошибка в periodic_sync: {e}")
            await asyncio.sleep(300)

async def verify_user_in_1c(full_name: str, employee_id: str) -> Optional[dict]:
    try:
        if not IC_CACHE["last_sync"] or datetime.now() - IC_CACHE["last_sync"] > timedelta(hours=1):
            await sync_with_1c()
        norm_id = normalize_employee_id(employee_id)
        norm_name = normalize_name(full_name)
        emp = EMPLOYEES_BY_NORM_ID.get(norm_id)
        if emp and emp.get('norm_name') == norm_name:
            return {
                "verified": True,
                "department": emp.get("department", ""),
                "position": emp.get("position", ""),
                "employee_id": emp.get("employee_id", employee_id),
                "from_cache": True
            }
        result = await verify_employee(full_name, employee_id)
        logger.info(f"Employee verification result for {full_name} (ID: {employee_id}): {result}")
        return result
    except Exception as e:
        logger.error(f"Employee verification error: {e}")
        return None

# ========== ФУНКЦИИ ДЛЯ РАБОТЫ С ДОКУМЕНТАМИ ==========

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            text = re.sub(r'\s+', ' ', para.text.strip())
            if text:
                full_text.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = re.sub(r'\s+', ' ', cell.text.strip())
                    if text:
                        full_text.append(text)
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Ошибка при чтении документа: {e}")
        return ""

def split_text_into_chunks(text: str, chunk_size: int = 500) -> List[str]:
    sentences = text.split('.')
    chunks = []
    current_chunk = []
    current_length = 0
    for sentence in sentences:
        sentence = sentence.strip() + '.'
        sentence_length = len(sentence.split())
        if current_length + sentence_length > chunk_size:
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_length = sentence_length
        else:
            current_chunk.append(sentence)
            current_length += sentence_length
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    return chunks

def extract_metadata_from_text(text: str) -> Dict[str, str]:
    """Извлекает метаданные из текста документа"""
    metadata = {}
    
    # Определяем тип документа
    text_lower = text.lower()
    if any(keyword in text_lower for keyword in ['регламент', 'процедура', 'инструкция']):
        metadata['doc_type'] = 'regulation'
    elif any(keyword in text_lower for keyword in ['faq', 'часто задаваемые', 'вопросы и ответы']):
        metadata['doc_type'] = 'faq'
    elif any(keyword in text_lower for keyword in ['справка', 'руководство', 'помощь']):
        metadata['doc_type'] = 'guide'
    else:
        metadata['doc_type'] = 'document'
    
    # Ищем упоминания отделов
    departments = ['hr', 'ит', 'бухгалтерия', 'продажи', 'маркетинг', 'администрация', 'безопасность']
    for dept in departments:
        if dept in text_lower:
            metadata['department'] = dept
            break
    
    # Ищем даты (простой паттерн)
    date_patterns = [
        r'\d{1,2}\.\d{1,2}\.\d{4}',  # 01.01.2024
        r'\d{4}-\d{1,2}-\d{1,2}',   # 2024-01-01
        r'\d{1,2}/\d{1,2}/\d{4}'    # 01/01/2024
    ]
    
    for pattern in date_patterns:
        matches = re.findall(pattern, text)
        if matches:
            metadata['dates'] = matches[:3]  # Максимум 3 даты
            break
    
    return metadata

def smart_chunk_documents(text: str, filename: str = "") -> List[Dict]:
    """Умное разбиение документов на чанки с метаданными"""
    chunks = []
    
    # Извлекаем общие метаданные документа
    doc_metadata = extract_metadata_from_text(text)
    doc_metadata['filename'] = filename
    
    # Определяем стратегию разбиения на основе типа документа
    doc_type = doc_metadata.get('doc_type', 'document')
    
    if doc_type == 'regulation':
        # Для регламентов - разбиваем по пунктам
        chunk_pattern = r'(?:^|\n)\s*\d+(?:\.\d+)*\.?\s+[А-ЯЁ].*?(?=(?:^|\n)\s*\d+(?:\.\d+)*\.?\s+[А-ЯЁ]|$)'
        min_chunk_size = 100
        max_chunk_size = 800
    elif doc_type == 'faq':
        # Для FAQ - разбиваем по вопросам
        chunk_pattern = r'(?:^|\n)\s*(?:В:|Вопрос:|Q:).*?(?=(?:^|\n)\s*(?:В:|Вопрос:|Q:)|$)'
        min_chunk_size = 50
        max_chunk_size = 600
    else:
        # Стандартное разбиение по абзацам и предложениям
        chunk_pattern = r'[^.!?]*[.!?]+(?:\s+[^.!?]*[.!?]+)*'
        min_chunk_size = 150
        max_chunk_size = 500
    
    # Пытаемся умное разбиение
    try:
        matches = re.finditer(chunk_pattern, text, re.DOTALL | re.MULTILINE)
        potential_chunks = [match.group().strip() for match in matches]
        
        # Фильтруем и объединяем слишком маленькие чанки
        current_chunk = ""
        chunk_counter = 0
        
        for chunk_text in potential_chunks:
            if not chunk_text:
                continue
                
            # Если чанк слишком маленький, объединяем с предыдущим
            if len(chunk_text) < min_chunk_size and current_chunk:
                current_chunk += " " + chunk_text
            else:
                # Сохраняем предыдущий чанк если он есть
                if current_chunk and len(current_chunk) >= min_chunk_size:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'chunk_id': chunk_counter,
                        'metadata': {**doc_metadata, 'chunk_size': len(current_chunk)}
                    })
                    chunk_counter += 1
                
                current_chunk = chunk_text
            
            # Если чанк слишком большой, разбиваем дальше
            if len(current_chunk) > max_chunk_size:
                # Простое разбиение по предложениям
                sentences = current_chunk.split('. ')
                temp_chunk = ""
                
                for sentence in sentences:
                    if len(temp_chunk + sentence) > max_chunk_size and temp_chunk:
                        chunks.append({
                            'text': temp_chunk.strip() + '.',
                            'chunk_id': chunk_counter,
                            'metadata': {**doc_metadata, 'chunk_size': len(temp_chunk)}
                        })
                        chunk_counter += 1
                        temp_chunk = sentence
                    else:
                        temp_chunk += sentence + ". " if temp_chunk else sentence
                
                current_chunk = temp_chunk
        
        # Добавляем последний чанк
        if current_chunk and len(current_chunk) >= min_chunk_size:
            chunks.append({
                'text': current_chunk.strip(),
                'chunk_id': chunk_counter,
                'metadata': {**doc_metadata, 'chunk_size': len(current_chunk)}
            })
        
    except Exception as e:
        logger.error(f"Ошибка умного разбиения: {e}, fallback к простому")
        # Fallback к простому разбиению
        simple_chunks = split_text_into_chunks(text, max_chunk_size)
        for i, chunk_text in enumerate(simple_chunks):
            chunks.append({
                'text': chunk_text,
                'chunk_id': i,
                'metadata': {**doc_metadata, 'chunk_size': len(chunk_text), 'method': 'simple'}
            })
    
    logger.info(f"Документ разбит на {len(chunks)} чанков (тип: {doc_type})")
    return chunks

async def index_document(file_path: str) -> bool:
    global document_embeddings, document_chunks
    try:
        text = extract_text_from_docx(file_path)
        if not text:
            return False
        chunks = split_text_into_chunks(text)
        embeddings = await llm_client.create_embeddings(chunks)
        if embeddings is None:
            return False
        embeddings = np.array(embeddings)
        document_chunks.extend(chunks)
        if document_embeddings is None:
            document_embeddings = embeddings
        else:
            document_embeddings = np.vstack([document_embeddings, embeddings])
        logger.info(f"Документ {file_path} успешно проиндексирован")
        return True
    except Exception as e:
        logger.error(f"Ошибка при индексации документа: {e}")
        return False

async def awaitable_create_query_embedding(query: str):
    try:
        return await llm_client.create_embeddings([query])
    except Exception:
        return None

def find_relevant_chunks(query: str, top_k: int = 3) -> List[str]:
    if not document_embeddings is None and len(document_chunks) > 0:
        try:
            loop = asyncio.get_event_loop()
            query_embedding_list = loop.run_until_complete(awaitable_create_query_embedding(query))
            if query_embedding_list is None:
                return []
            query_embedding = np.array(query_embedding_list[0])
            similarities = np.dot(document_embeddings, query_embedding)
            top_indices = np.argsort(similarities)[-top_k:][::-1]
            return [document_chunks[i] for i in top_indices]
        except Exception as e:
            logger.error(f"Ошибка при поиске релевантных чанков: {e}")
    return []

async def generate_response(query: str, context: str = "") -> str:
    try:
        response = await llm_client.generate(query=query, context=context, max_tokens=MAX_NEW_TOKENS)
        if response is None:
            return "Извините, произошла ошибка при обработке вашего запроса."
        return response
    except Exception as e:
        logger.error(f"Ошибка при генерации ответа: {e}")
        return "Извините, произошла ошибка при обработке вашего запроса."

# Глобальные переменные для RAG
document_embeddings = None
document_chunks = []

# Для хранения сессий Q&A и связывания с feedback
QA_SESSIONS: Dict[int, int] = {}  # message_id -> qa_session_id

# A/B тестирование и конфигурация
USE_SEARCH_V2 = os.getenv('USE_SEARCH_V2', 'false').lower() == 'true'
SEARCH_V2_PERCENTAGE = int(os.getenv('SEARCH_V2_PERCENTAGE', '30'))  # % пользователей на новой версии

# Дополнительные вспомогательные для индексации сервиса
ALLOWED_EXTENSIONS = {'.docx'}

def _sanitize_filename(name: str) -> str:
    name = re.sub(r"[^\w\-\.]+", "_", name)
    return name.strip("._") or f"doc_{int(datetime.now().timestamp())}.docx"

async def expand_query(original_query: str) -> List[str]:
    """Расширение запроса: генерация альтернативных формулировок"""
    try:
        # Генерируем перефразировки
        rephrase_prompt = f"""Перефразируй этот вопрос 2 способами для лучшего поиска в корпоративной документации:

Оригинальный вопрос: "{original_query}"

Варианты:
1."""
        
        response = await llm_client.generate(
            query=rephrase_prompt,
            max_tokens=150,
            temperature=0.3
        )
        
        if response:
            # Извлекаем варианты из ответа
            variants = []
            for line in response.split('\n'):
                line = line.strip()
                if line and any(line.startswith(prefix) for prefix in ['1.', '2.', '•', '-']):
                    # Убираем нумерацию
                    clean_line = re.sub(r'^[12]\.\s*|^[•-]\s*', '', line).strip()
                    if clean_line and len(clean_line) > 10:
                        variants.append(clean_line)
            
            # Возвращаем оригинал + максимум 2 варианта
            result = [original_query] + variants[:2]
            logger.info(f"Query expansion: {len(result)} variants generated")
            return result
        
    except Exception as e:
        logger.error(f"Ошибка при расширении запроса: {e}")
    
    # Fallback - возвращаем только оригинальный запрос
    return [original_query]

def should_use_search_v2(user_id: int) -> bool:
    """Определяет, использовать ли новую версию поиска для пользователя"""
    if USE_SEARCH_V2:
        return True
    
    # A/B тест: определённый процент пользователей на новой версии
    return (user_id % 100) < SEARCH_V2_PERCENTAGE

async def search_documents(query: str, user_id: int = 0, use_expansion: bool = True) -> tuple:
    """Универсальная функция поиска с выбором версии и расширением запросов"""
    try:
        search_version = "v2" if should_use_search_v2(user_id) else "v1"
        endpoint = "/search_v2" if search_version == "v2" else "/search"
        
        # Query expansion для улучшения результатов
        queries_to_search = [query]
        if use_expansion and len(query) > 20:  # Расширяем только длинные запросы
            queries_to_search = await expand_query(query)
        
        all_hits = []
        confidence_scores = []
        
        import aiohttp
        async with aiohttp.ClientSession() as session:
            for search_query in queries_to_search:
                async with session.post(f"{llm_client.base_url}{endpoint}", 
                                      json={"query": search_query, "top_k": 3}) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        hits = data.get('hits', [])
                        if hits:
                            all_hits.extend(hits)
                            confidence_scores.extend([h['score'] for h in hits])
        
        # Дедупликация по тексту и ранжирование
        seen_texts = set()
        unique_hits = []
        for hit in all_hits:
            text_hash = hash(hit['text'][:100])  # Хеш первых 100 символов
            if text_hash not in seen_texts:
                seen_texts.add(text_hash)
                unique_hits.append(hit)
        
        # Сортируем по скору и берём топ-3
        unique_hits.sort(key=lambda x: x['score'], reverse=True)
        final_hits = unique_hits[:3]
        
        if final_hits:
            top_context = "\n\n".join(h['text'] for h in final_hits)
            sources_block = f"\n\nИсточники ({search_version}):\n" + "\n".join([f"— {round(h['score'],3)}" for h in final_hits])
            top_score = final_hits[0]['score']
            
            return top_context, sources_block, top_score, True, search_version
        else:
            return "", "", 0.0, False, search_version
            
    except Exception as e:
        logger.error(f"Ошибка в search_documents: {e}")
        return "", "", 0.0, False, "error"

async def rebuild_service_index_from_docs() -> int:
    try:
        documents: List[str] = []
        for fname in os.listdir(DOCUMENTS_DIR):
            if not any(fname.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
                continue
            fpath = os.path.join(DOCUMENTS_DIR, fname)
            text = extract_text_from_docx(fpath)
            if not text:
                continue
            # Используем умное разбиение с метаданными
            smart_chunks = smart_chunk_documents(text, fname)
            # Извлекаем только текст для индексации (пока что)
            chunk_texts = [chunk['text'] for chunk in smart_chunks if chunk['text'].strip()]
            documents.extend(chunk_texts)
            
            logger.info(f"Обработан {fname}: {len(chunk_texts)} чанков")
            
        if not documents:
            logger.info("Нет документов для индексации")
            return 0
            
        import aiohttp
        async with aiohttp.ClientSession() as session:
            async with session.post(f"{llm_client.base_url}/index", json={"documents": documents}) as resp:
                if resp.status != 200:
                    err = await resp.text()
                    logger.error(f"Ошибка при обращении к /index: {err}")
                    return 0
                data = await resp.json()
                logger.info(f"Проиндексировано документов (чанков): {data}")
                
        return len(documents)
    except Exception as e:
        logger.error(f"Ошибка при перестроении индекса сервиса: {e}")
        return 0

def create_feedback_keyboard(qa_session_id: int) -> InlineKeyboardMarkup:
    """Создаёт клавиатуру с кнопками лайк/дизлайк для фидбека."""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="👍", callback_data=f"feedback_like_{qa_session_id}"),
            InlineKeyboardButton(text="👎", callback_data=f"feedback_dislike_{qa_session_id}")
        ]
    ])
    return keyboard

# ========== ОБРАБОТЧИКИ КОМАНД ==========

def setup_handlers(dp: Dispatcher):
    @dp.message(Command('start'))
    async def start_handler(message: types.Message):
        user_id = message.from_user.id
        if not API_TOKEN:
            await message.answer('❌ Не задан API_TOKEN. Укажите его в .env')
            return
        if user_id in AUTHORIZED_USERS:
            await message.answer('Добро пожаловать! Вы авторизованы.', reply_markup=main_kb)
            return
        if await can_try_registration(user_id):
            await set_user_state(user_id, 'step', 'name')
            await message.answer(
                'Для регистрации введите ваше ФИО\n'
                '(как указано в 1С, например: Иванов Иван Иванович)'
            )
        else:
            next_attempt = await get_next_attempt_time(user_id)
            await message.answer(
                f"⚠️ Превышено количество попыток регистрации.\n"
                f"Следующая попытка будет доступна {next_attempt}.\n\n"
                f"Если вам нужна помощь, обратитесь к администратору.",
                reply_markup=ReplyKeyboardRemove()
            )

    @dp.message(Command('cancel'))
    async def cancel_handler(message: types.Message):
        user_id = message.from_user.id
        await clear_user_state(user_id)
        await message.answer('Операция отменена. Используйте /start для регистрации.', reply_markup=ReplyKeyboardRemove())

    # Русская кнопка «Отмена»
    @dp.message(lambda m: m.text and m.text.strip().lower() == 'отмена')
    async def cancel_button(message: types.Message):
        await cancel_handler(message)

    @dp.message(Command('status'))
    async def status_handler(message: types.Message):
        user_id = message.from_user.id
        if user_id not in AUTHORIZED_USERS:
            await message.answer('❌ Недоступно до завершения регистрации. Используйте /start.')
            return
        info = AUTHORIZED_INFO.get(user_id, {})
        await message.answer(
            f"👤 ФИО: {info.get('name','-')}\n"
            f"🔢 Табельный: {info.get('employee_id','-')}\n"
            f"🕒 Регистрация: {info.get('verified_at','-')}",
            reply_markup=main_kb
        )

    # ====== Админ: добавление документов (SOP .docx) ======
    @dp.message(Command('train'))
    async def train_handler(message: types.Message):
        if message.from_user.id != ADMIN_CHAT_ID:
            await message.answer('❌ Команда доступна только администратору.')
            return
        await set_user_state(message.from_user.id, 'awaiting_doc_upload', '1')
        await message.answer(
            'Пришлите файл .docx (SOP/регламент). Я сохраню его и переиндексирую поиск.',
            reply_markup=ReplyKeyboardRemove()
        )

    @dp.message(lambda m: m.document is not None)
    async def handle_document_upload(message: types.Message):
        try:
            user_id = message.from_user.id
            doc = message.document
            file_name = (doc.file_name or 'document.docx')
            ext = os.path.splitext(file_name)[1].lower()
            awaiting = await get_user_state(user_id, 'awaiting_doc_upload')
            # Разрешаем только админу и только .docx
            if user_id != ADMIN_CHAT_ID:
                return
            if ext not in ALLOWED_EXTENSIONS:
                if awaiting:
                    await message.answer('❌ Поддерживаются только файлы .docx. Попробуйте снова.')
                return
            # Путь сохранения
            safe_name = _sanitize_filename(file_name)
            save_path = os.path.join(DOCUMENTS_DIR, safe_name)
            # Скачиваем файл
            await bot.download(doc, destination=save_path)
            await set_user_state(user_id, 'awaiting_doc_upload', '0')
            # Переиндексация всех документов
            chunks_count = await rebuild_service_index_from_docs()
            await message.answer(
                f"✅ Файл сохранён: {safe_name}\n📚 Индекс обновлён, чанков: {chunks_count}",
                reply_markup=main_kb
            )
        except Exception as e:
            logger.error(f"Ошибка загрузки документа: {e}")
            await message.answer('❌ Не удалось обработать документ. Проверьте формат и попробуйте снова.', reply_markup=main_kb)

    @dp.message(lambda message: message.text and message.text.lower() in RETRY_COMMANDS)
    async def retry_registration(message: types.Message):
        user_id = message.from_user.id
        if user_id in AUTHORIZED_USERS:
            await message.answer('Вы уже зарегистрированы.', reply_markup=main_kb)
            return
        if await can_try_registration(user_id):
            await clear_user_state(user_id)
            await set_user_state(user_id, 'step', 'name')
            await message.answer(
                "🔄 Начинаем регистрацию заново!\n\n"
                "Введите ваше ФИО\n"
                "(как указано в 1С, например: Иванов Иван Иванович)",
                reply_markup=ReplyKeyboardRemove()
            )
        else:
            next_attempt = await get_next_attempt_time(user_id)
            await message.answer(
                f"⚠️ Превышено количество попыток регистрации.\n"
                f"Следующая попытка будет доступна {next_attempt}.\n\n"
                f"Если вам нужна помощь, обратитесь к администратору.",
                reply_markup=ReplyKeyboardRemove()
            )

    @dp.message(lambda message: True)
    async def registration_handler(message: types.Message):
        user_id = message.from_user.id
        step = await get_user_state(user_id, 'step')
        if not step:
            return
        if step == 'name':
            name_parts = message.text.strip().split()
            if len(name_parts) != 3:
                await message.answer(
                    '❌ Пожалуйста, введите ФИО полностью\n'
                    'Формат: Фамилия Имя Отчество\n'
                    'Пример: Иванов Иван Иванович\n\n'
                    'Для повторной попытки напишите "повтор"'
                )
                return
            await set_user_state(user_id, 'name', message.text.strip())
            await set_user_state(user_id, 'step', 'employee_id')
            await message.answer(
                'Введите ваш табельный номер\n'
                '(как указано в 1С, например: E001)'
            )
        elif step == 'employee_id':
            employee_id = message.text.strip()
            name = await get_user_state(user_id, 'name')
            await set_user_state(user_id, 'employee_id', employee_id)
            verification_result = await verify_user_in_1c(name, employee_id)
            if verification_result and verification_result.get('verified'):
                await clear_user_state(user_id)
                AUTHORIZED_USERS.add(user_id)
                AUTHORIZED_INFO[user_id] = {
                    'name': name,
                    'employee_id': employee_id,
                    'department': verification_result.get('department',''),
                    'position': verification_result.get('position',''),
                    'verified_at': datetime.now().strftime('%Y-%m-%d %H:%M')
                }
                await log_registration_attempt(user_id, name, employee_id, True)
                await message.answer(
                    f"✅ Регистрация подтверждена!\n\n"
                    f"📋 Ваши данные:\n"
                    f"👤 ФИО: {name}\n"
                    f"🔢 Табельный номер: {employee_id}\n"
                    f"📝 Должность: {verification_result.get('position', 'Не указана')}\n"
                    f"🏢 Отдел: {verification_result.get('department', 'Не указан')}\n\n"
                    f"Теперь вам доступны команды /ask и /help.",
                    reply_markup=main_kb
                )
            else:
                await inc_registration_attempt(user_id)
                await log_registration_attempt(user_id, name or '', employee_id, False)
                attempts = REGISTRATION_ATTEMPTS.get(user_id, {}).get('count', 0)
                attempts_left = MAX_ATTEMPTS - attempts
                error_msg = [
                    "❌ Ошибка проверки данных\n",
                    "Возможные причины:",
                    "1. Неверно указано ФИО",
                    "2. Неверный табельный номер",
                    "3. Данные не совпадают с базой\n"
                ]
                if attempts_left > 0:
                    error_msg.append(f"У вас осталось попыток: {attempts_left}")
                    error_msg.append("Для повторной попытки нажмите кнопку 'повтор' или /cancel")
                else:
                    next_attempt = await get_next_attempt_time(user_id)
                    error_msg.append(f"Следующая попытка будет доступна {next_attempt}")
                await message.answer('\n'.join(error_msg), reply_markup=retry_kb)
                await clear_user_state(user_id)

    # Команда /ask
    @dp.message(Command('ask'))
    async def ask_handler(message: types.Message):
        user_id = message.from_user.id
        if user_id not in AUTHORIZED_USERS:
            await message.answer("❌ Вы не авторизованы! Используйте /start для регистрации.")
            return
        awaiting = await get_user_state(user_id, 'awaiting_question')
        if not awaiting:
            await set_user_state(user_id, 'awaiting_question', '1')
        await message.answer(
            "Задайте ваш вопрос. Я постараюсь ответить, используя доступную документацию.",
            reply_markup=ReplyKeyboardRemove()
        )

    # Кнопка «Спросить» (русская)
    @dp.message(lambda m: m.text and m.text.strip().lower() == 'спросить')
    async def ask_button(message: types.Message):
        await ask_handler(message)

    @dp.message(lambda message: True)
    async def process_question(message: types.Message):
        user_id = message.from_user.id
        if user_id not in AUTHORIZED_USERS:
            return
        awaiting = await get_user_state(user_id, 'awaiting_question')
        if not awaiting:
            return
        await set_user_state(user_id, 'awaiting_question', '0')
        
        start_time = time.time()
        processing_msg = await message.answer("🤔 Обрабатываю ваш вопрос...")
        
        try:
            import aiohttp
            top_context = ""
            sources_block = ""
            conf_threshold = 0.12
            confidence_score = 0.0
            context_found = False
            
            try:
                # Используем универсальную функцию поиска
                top_context, sources_block, confidence_score, context_found, search_version = await search_documents(message.text, user_id)
                
                if not context_found:
                    # Логируем как неотвеченный вопрос
                    user_info = AUTHORIZED_INFO.get(user_id, {})
                    await log_unanswered_question(
                        user_id, 
                        message.text, 
                        f"Department: {user_info.get('department', '')}, Position: {user_info.get('position', '')}"
                    )
                    await processing_msg.edit_text(
                        "Я не уверен в ответе. Уточните вопрос или добавьте деталей (дата, подразделение, документ)."
                    )
                    return
                
                if confidence_score < conf_threshold:
                    # Логируем как неотвеченный вопрос
                    user_info = AUTHORIZED_INFO.get(user_id, {})
                    await log_unanswered_question(
                        user_id, 
                        message.text, 
                        f"Department: {user_info.get('department', '')}, Position: {user_info.get('position', '')}, Search: {search_version}"
                    )
                    await processing_msg.edit_text(
                        f"Я не уверен в ответе (поиск: {search_version}). Уточните вопрос или добавьте деталей (дата, подразделение, документ)."
                    )
                    return
            except Exception as e:
                logger.error(f"Ошибка поиска контекста: {e}")

            context = top_context
            response = await generate_response(message.text, context)
            
            # Замеряем время ответа
            response_time_ms = int((time.time() - start_time) * 1000)
            
            # Логируем сессию Q&A с информацией о версии поиска
            user_info = AUTHORIZED_INFO.get(user_id, {})
            qa_session_id = await log_qa_session(
                telegram_id=user_id,
                user_name=user_info.get('name', message.from_user.full_name or ''),
                employee_id=user_info.get('employee_id', ''),
                question=message.text,
                answer=f"[{search_version}] {response}",  # Помечаем версию поиска
                response_time_ms=response_time_ms,
                confidence_score=confidence_score,
                context_found=context_found
            )
            
            # Отправляем ответ с кнопками фидбека
            feedback_kb = create_feedback_keyboard(qa_session_id)
            sent_msg = await processing_msg.edit_text(
                f"Вопрос: {message.text}\n\n"
                f"Ответ: {response}{sources_block}",
                reply_markup=feedback_kb
            )
            
            # Сохраняем связь сообщения и сессии для обработки фидбека
            if qa_session_id:
                QA_SESSIONS[sent_msg.message_id] = qa_session_id
                
        except Exception as e:
            logger.error(f"Ошибка при обработке вопроса: {e}")
            await processing_msg.edit_text(
                "😔 Извините, произошла ошибка при обработке вашего вопроса. Попробуйте позже.",
                reply_markup=main_kb
            )

    # Обработчик callback для кнопок фидбека
    @dp.callback_query(lambda c: c.data and c.data.startswith('feedback_'))
    async def process_feedback(callback_query: types.CallbackQuery):
        try:
            data_parts = callback_query.data.split('_')
            if len(data_parts) != 3:
                return
                
            action = data_parts[1]  # 'like' или 'dislike'
            qa_session_id = int(data_parts[2])
            
            rating = 1 if action == 'like' else -1
            success = await save_feedback(qa_session_id, callback_query.from_user.id, rating)
            
            if success:
                emoji = "👍" if rating == 1 else "👎"
                await callback_query.answer(f"Спасибо за оценку! {emoji}")
                
                # Убираем кнопки после голосования
                await callback_query.message.edit_reply_markup(reply_markup=None)
            else:
                await callback_query.answer("Ошибка при сохранении оценки", show_alert=True)
                
        except Exception as e:
            logger.error(f"Ошибка при обработке фидбека: {e}")
            await callback_query.answer("Произошла ошибка", show_alert=True)

    # Команда /help
    @dp.message(Command('help'))
    async def help_handler(message: types.Message):
        if message.from_user.id not in AUTHORIZED_USERS:
            await message.answer("❌ Недоступно до завершения регистрации. Используйте /start.")
            return
        help_text = (
            '<b>Доступные команды:</b>\n\n'
            '/start - начать регистрацию\n'
            '/status - показать профиль\n'
            '/ask - задать вопрос ассистенту\n'
            '/cancel - отменить текущую операцию\n'
            "'повтор' - повторить попытку регистрации\n"
            '/help - помощь'
        )
        if message.from_user.id == ADMIN_CHAT_ID:
            help_text += '\n\n<b>Админ-команды:</b>\n' \
                        '/train - добавить документ (.docx) и обновить поиск\n' \
                        '/analytics - подробная статистика использования\n' \
                        '/stats - краткая статистика (алиас для analytics)\n' \
                        '/compare_search - сравнение версий поиска (A/B тест)'
        await message.answer(help_text, parse_mode='HTML', reply_markup=main_kb)

    # Кнопка «Помощь» (русская)
    @dp.message(lambda m: m.text and m.text.strip().lower() == 'помощь')
    async def help_button(message: types.Message):
        await help_handler(message)

    # ====== Админ: аналитика и статистика ======
    @dp.message(Command('analytics'))
    async def analytics_handler(message: types.Message):
        if message.from_user.id != ADMIN_CHAT_ID:
            await message.answer('❌ Команда доступна только администратору.')
            return
        
        try:
            # Получаем статистику за последние 7 дней
            stats = await get_analytics_stats(days=7)
            popular_questions = await get_popular_questions(limit=5, days=30)
            
            if not stats:
                await message.answer('📊 Пока нет данных для аналитики.')
                return
            
            report_lines = [
                '📊 <b>Аналитика бота за 7 дней</b>\n',
                f'📈 <b>Общая статистика:</b>',
                f'• Всего вопросов: {stats.get("total_questions", 0)}',
                f'• Среднее время ответа: {stats.get("avg_response_time_ms", 0)} мс',
                f'• Средняя уверенность: {stats.get("avg_confidence", 0)}',
                f'• Вопросов с контекстом: {stats.get("questions_with_context", 0)}',
                '',
                f'👥 <b>Удовлетворённость пользователей:</b>',
                f'• Лайков: {stats.get("likes", 0)} 👍',
                f'• Дизлайков: {stats.get("dislikes", 0)} 👎',
                f'• Уровень удовлетворённости: {stats.get("satisfaction_rate", 0)}%',
                ''
            ]
            
            # Топ вопросов за 7 дней
            if stats.get('top_questions'):
                report_lines.append('🔥 <b>Популярные вопросы (7 дней):</b>')
                for i, (question, freq) in enumerate(stats['top_questions'][:5], 1):
                    short_q = question[:50] + '...' if len(question) > 50 else question
                    report_lines.append(f'{i}. "{short_q}" ({freq}x)')
                report_lines.append('')
            
            # Неотвеченные вопросы
            if stats.get('unanswered_questions'):
                report_lines.append('❓ <b>Неотвеченные вопросы:</b>')
                for question, freq in stats['unanswered_questions'][:5]:
                    short_q = question[:50] + '...' if len(question) > 50 else question
                    report_lines.append(f'• "{short_q}" ({freq}x)')
                report_lines.append('')
            
            # Популярные вопросы за месяц
            if popular_questions:
                report_lines.append('📅 <b>Топ вопросов за месяц:</b>')
                for i, (question, freq, avg_conf) in enumerate(popular_questions[:3], 1):
                    short_q = question[:40] + '...' if len(question) > 40 else question
                    conf_str = f'(уверенность: {round(avg_conf or 0, 2)})' if avg_conf else ''
                    report_lines.append(f'{i}. "{short_q}" ({freq}x) {conf_str}')
            
            report_text = '\n'.join(report_lines)
            await message.answer(report_text, parse_mode='HTML')
            
        except Exception as e:
            logger.error(f"Ошибка при формировании аналитики: {e}")
            await message.answer('❌ Ошибка при получении статистики. Проверьте логи.')

    @dp.message(Command('stats'))
    async def stats_handler(message: types.Message):
        """Краткая статистика (алиас для analytics)"""
        if message.from_user.id != ADMIN_CHAT_ID:
            await message.answer('❌ Команда доступна только администратору.')
            return
        await analytics_handler(message)

    # ====== Админ: сравнение версий поиска ======
    @dp.message(Command('compare_search'))
    async def compare_search_handler(message: types.Message):
        if message.from_user.id != ADMIN_CHAT_ID:
            await message.answer('❌ Команда доступна только администратору.')
            return
        
        try:
            # Анализируем логи за последние 7 дней
            import aiosqlite
            async with aiosqlite.connect('employees.db') as db:
                # Статистика v1 (без префикса [v2])
                async with db.execute(
                    """SELECT 
                        COUNT(*) as total,
                        AVG(response_time_ms) as avg_time,
                        AVG(confidence_score) as avg_confidence
                    FROM qa_sessions 
                    WHERE created_at >= datetime('now', '-7 days')
                    AND (answer NOT LIKE '[v2]%' OR answer IS NULL)"""
                ) as cursor:
                    v1_stats = await cursor.fetchone()
                
                # Статистика v2 (с префиксом [v2])
                async with db.execute(
                    """SELECT 
                        COUNT(*) as total,
                        AVG(response_time_ms) as avg_time,
                        AVG(confidence_score) as avg_confidence
                    FROM qa_sessions 
                    WHERE created_at >= datetime('now', '-7 days')
                    AND answer LIKE '[v2]%'"""
                ) as cursor:
                    v2_stats = await cursor.fetchone()
                
                # Фидбек по версиям
                async with db.execute(
                    """SELECT 
                        CASE WHEN qa.answer LIKE '[v2]%' THEN 'v2' ELSE 'v1' END as version,
                        SUM(CASE WHEN f.rating = 1 THEN 1 ELSE 0 END) as likes,
                        SUM(CASE WHEN f.rating = -1 THEN 1 ELSE 0 END) as dislikes
                    FROM feedback f
                    JOIN qa_sessions qa ON f.qa_session_id = qa.id
                    WHERE qa.created_at >= datetime('now', '-7 days')
                    GROUP BY version"""
                ) as cursor:
                    feedback_stats = await cursor.fetchall()
            
            # Формируем отчёт
            report_lines = [
                '🔍 <b>Сравнение версий поиска (7 дней)</b>\n',
                f'📊 <b>Конфигурация:</b>',
                f'• USE_SEARCH_V2: {USE_SEARCH_V2}',
                f'• SEARCH_V2_PERCENTAGE: {SEARCH_V2_PERCENTAGE}%',
                '',
                f'🔧 <b>Поиск V1 (классический):</b>',
                f'• Запросов: {v1_stats[0] if v1_stats else 0}',
                f'• Среднее время: {round(v1_stats[1] or 0)} мс',
                f'• Средняя уверенность: {round(v1_stats[2] or 0, 3)}',
                '',
                f'⚡ <b>Поиск V2 (Cross-Encoder):</b>',
                f'• Запросов: {v2_stats[0] if v2_stats else 0}',
                f'• Среднее время: {round(v2_stats[1] or 0)} мс',
                f'• Средняя уверенность: {round(v2_stats[2] or 0, 3)}',
                ''
            ]
            
            # Добавляем сравнение фидбека
            if feedback_stats:
                report_lines.append('👥 <b>Удовлетворённость пользователей:</b>')
                for version, likes, dislikes in feedback_stats:
                    total = likes + dislikes
                    satisfaction = round(likes / total * 100, 1) if total > 0 else 0
                    report_lines.append(f'• {version.upper()}: {satisfaction}% ({likes}👍 / {dislikes}👎)')
                report_lines.append('')
            
            # Рекомендации
            if v2_stats and v1_stats and v2_stats[0] > 0 and v1_stats[0] > 0:
                v2_faster = (v2_stats[1] or 0) < (v1_stats[1] or 0)
                v2_more_confident = (v2_stats[2] or 0) > (v1_stats[2] or 0)
                
                report_lines.append('💡 <b>Рекомендации:</b>')
                if v2_faster and v2_more_confident:
                    report_lines.append('✅ V2 превосходит V1 - рекомендуется увеличить SEARCH_V2_PERCENTAGE')
                elif v2_more_confident:
                    report_lines.append('✅ V2 точнее, но медленнее - можно увеличить долю постепенно')
                else:
                    report_lines.append('⚠️ V1 пока показывает лучшие результаты')
            
            await message.answer('\n'.join(report_lines), parse_mode='HTML')
            
        except Exception as e:
            logger.error(f"Ошибка при сравнении версий поиска: {e}")
            await message.answer('❌ Ошибка при получении статистики сравнения.')

        if message.from_user.id == ADMIN_CHAT_ID:
            help_text += '\n\n<b>Админ-команды:</b>\n' \
                        '/train - добавить документ (.docx) и обновить поиск\n' \
                        '/analytics - подробная статистика использования\n' \
                        '/stats - краткая статистика (алиас для analytics)\n' \
                        '/compare_search - сравнение версий поиска (A/B тест)'